import os
import re
import random
from datetime import datetime, timedelta
import PyPDF2
from docx import Document
import requests
from bs4 import BeautifulSoup
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

# Ensure NLTK data is downloaded
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')

def extract_text_from_file(filepath):
    """Extract text from PDF, DOCX, or TXT files with error handling"""
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    text = ""
    try:
        if filepath.lower().endswith('.pdf'):
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if not reader.pages:
                    raise ValueError("PDF contains no readable pages")
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif filepath.lower().endswith('.docx'):
            doc = Document(filepath)
            if not doc.paragraphs:
                raise ValueError("DOCX contains no readable text")
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filepath.lower().endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        raise ValueError(f"Error reading file: {str(e)}")
    
    return text.strip()

def extract_text_from_url(url):
    """Extract text content from a URL with error handling"""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove unwanted elements
        for element in soup(['script', 'style', 'nav', 'footer', 'iframe']):
            element.decompose()
            
        # Get text from paragraphs and headers
        text = ' '.join([p.get_text(' ', strip=True) 
                        for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4'])])
        
        return re.sub(r'\s+', ' ', text).strip()
    except Exception as e:
        raise ValueError(f"Error extracting URL content: {str(e)}")

def validate_dates(start_date, end_date):
    """Validate date range and format"""
    try:
        start = datetime.strptime(start_date, '%Y-%m-%d')
        end = datetime.strptime(end_date, '%Y-%m-%d')
        
        if end < start:
            raise ValueError("End date must be after start date")
        if (end - start).days > 365:
            raise ValueError("Study period cannot exceed 1 year")
            
        return start, end
    except ValueError as e:
        raise ValueError(f"Invalid dates: {str(e)}")

def process_document(filepath, start_date, end_date, daily_hours):
    """Main document processing function with comprehensive error handling"""
    try:
        # Validate inputs
        if not filepath:
            raise ValueError("No filepath or URL provided")
        
        start, end = validate_dates(start_date, end_date)
        total_days = (end - start).days
        
        if daily_hours <= 0 or daily_hours > 12:
            raise ValueError("Daily study hours must be between 0.5 and 12")
        
        # Extract text content
        if filepath.startswith(('http://', 'https://')):
            text = extract_text_from_url(filepath)
        else:
            text = extract_text_from_file(filepath)
        
        if not text:
            raise ValueError("Document contains no readable text")
        
        # Basic text cleaning
        text = re.sub(r'\s+', ' ', text).strip()
        sentences = [s.strip() for s in sent_tokenize(text) if s.strip()]
        
        if not sentences:
            raise ValueError("Could not extract meaningful sentences")
        
        # Calculate study chunks
        available_hours = total_days * daily_hours
        target_chunks = max(1, min(len(sentences), int(available_hours * 2)))  # 30-min chunks
        
        # Distribute sentences to chunks
        chunks = []
        sentences_per_chunk = max(1, len(sentences) // target_chunks)
        
        for i in range(0, len(sentences), sentences_per_chunk):
            chunk_sentences = sentences[i:i+sentences_per_chunk]
            chunks.append({
                'content': ' '.join(chunk_sentences),
                'duration': 0.5  # 30 minutes per chunk
            })
        
        # Generate study plan
        study_plan = []
        current_date = start
        for chunk in chunks:
            study_plan.append({
                'date': current_date.strftime('%Y-%m-%d'),
                'content': chunk['content'],
                'duration_hours': chunk['duration']
            })
            current_date += timedelta(days=1)
        
        return {
            'word_count': len(word_tokenize(text)),
            'estimated_hours': len(sentences) / 120,  # Approx 2 sentences per minute
            'available_hours': available_hours,
            'study_plan': study_plan,
            'content_summary': summarize_content(sentences)
        }
        
    except Exception as e:
        raise ValueError(f"Document processing failed: {str(e)}")

def summarize_content(sentences):
    """Generate a simple summary from sentences"""
    try:
        if len(sentences) <= 3:
            return ' '.join(sentences)
        
        # Simple summary: first, last, and a middle sentence
        return '\n'.join([
            sentences[0],
            sentences[len(sentences)//2],
            sentences[-1]
        ])
    except:
        return "Could not generate summary"

def answer_question(filepath, question):
    """Basic question answering using keyword matching"""
    try:
        # Extract text
        if filepath.startswith(('http://', 'https://')):
            text = extract_text_from_url(filepath)
        else:
            text = extract_text_from_file(filepath)
        
        if not text:
            return "No content available to answer the question."
        
        # Simple keyword matching
        sentences = sent_tokenize(text)
        question_words = set(word.lower() for word in word_tokenize(question) 
                           if word.isalpha() and len(word) > 2)
        
        best_match = ""
        best_score = 0
        
        for sentence in sentences:
            sentence_words = set(word.lower() for word in word_tokenize(sentence)
                               if word.isalpha() and len(word) > 2)
            common_words = question_words.intersection(sentence_words)
            score = len(common_words)
            
            if score > best_score:
                best_score = score
                best_match = sentence
        
        return best_match if best_match else "Answer not found in the material."
    
    except Exception as e:
        return f"Error answering question: {str(e)}"

def generate_quiz(filepath, difficulty='medium'):
    """Generate quiz questions from document content"""
    try:
        # Extract text
        if filepath.startswith(('http://', 'https://')):
            text = extract_text_from_url(filepath)
        else:
            text = extract_text_from_file(filepath)
        
        if not text:
            return []
        
        sentences = [s for s in sent_tokenize(text) if len(word_tokenize(s)) > 6]
        
        # Determine question count based on difficulty
        question_counts = {'easy': 5, 'medium': 10, 'hard': 15, 'pro': 20}
        num_questions = min(question_counts.get(difficulty, 10), len(sentences))
        
        questions = []
        used_sentences = set()
        
        while len(questions) < num_questions and len(used_sentences) < len(sentences):
            # Find an unused sentence
            sentence = None
            for s in sentences:
                if s not in used_sentences and len(word_tokenize(s)) > 8:
                    sentence = s
                    used_sentences.add(s)
                    break
            
            if not sentence:
                break
            
            # Create fill-in-the-blank question
            words = word_tokenize(sentence)
            blank_pos = random.randint(1, len(words)-2)
            correct = words[blank_pos]
            
            # Skip if blank word is too short or not alphabetic
            if len(correct) <= 2 or not correct.isalpha():
                continue
            
            # Generate question and options
            question_text = ' '.join(words[:blank_pos] + ['______'] + words[blank_pos+1:])
            options = [correct]
            
            # Add similar words as distractors
            for word in word_tokenize(text):
                if (word[0].lower() == correct[0].lower() and 
                    word.lower() != correct.lower() and 
                    word.isalpha() and
                    len(word) > 2 and
                    word not in options):
                    options.append(word)
                    if len(options) >= 4:
                        break
            
            # Fill with random words if needed
            while len(options) < 4:
                random_word = random.choice(word_tokenize(text))
                if random_word.isalpha() and random_word not in options:
                    options.append(random_word)
            
            random.shuffle(options)
            correct_index = options.index(correct)
            
            questions.append({
                'question': f"Fill in the blank: {question_text}",
                'options': options,
                'correct_answer': chr(65 + correct_index)  # A, B, C, or D
            })
        
        return questions
    
    except Exception as e:
        print(f"Error generating quiz: {str(e)}")
        return []